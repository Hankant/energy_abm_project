from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\69596\.gemini\antigravity\scratch\energy_abm_project")
SOURCE = ROOT / "docs" / "planning" / "policy_mix_survey_abm_methodological_report.md"
OUTPUT = ROOT / "output" / "doc" / "policy_mix_survey_abm_methodological_report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_inline_runs(paragraph, text):
    pattern = re.compile(r"(\*\*.*?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`.*?`|https?://\S+)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(token.rstrip(".,;"))
            run.font.color.rgb = RGBColor(5, 99, 193)
            run.underline = True
            if token[-1:] in ".,;":
                paragraph.add_run(token[-1])
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_paragraph(doc, text, style=None, indent=0):
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    add_inline_runs(p, text)
    return p


def add_table(doc, rows):
    widths = len(rows[0])
    table = doc.add_table(rows=1, cols=widths)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, value in enumerate(rows[0]):
        cell = hdr.cells[i]
        cell.text = value.strip()
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    set_repeat_table_header(hdr)
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = value.strip()
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for name, size, color in [
        ("Title", 20, "17365D"),
        ("Heading 1", 15, "17365D"),
        ("Heading 2", 12.5, "1F4E79"),
        ("Heading 3", 11, "365F91"),
    ]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    # title page
    title = lines[0].lstrip("# ")
    subtitle = "面向导师讨论的研究论证报告（精简修订版）"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(23, 54, 93)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(24)
    r2 = p2.add_run(subtitle)
    r2.font.name = "楷体"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    r2.font.size = Pt(15)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(210)
    p3.add_run("能源政策组合与居民行为仿真项目\n2026年7月").font.size = Pt(11)
    doc.add_page_break()

    i = 1
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
            run.font.size = Pt(9.5)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F3F6F8")
            p_pr.append(shd)
            i += 1
            continue
        if line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for idx, tline in enumerate(table_lines):
                cells = [c.strip() for c in tline.strip("|").split("|")]
                if idx == 1 and all(set(c) <= set("-: ") for c in cells):
                    continue
                rows.append(cells)
            if rows:
                add_table(doc, rows)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("> "):
            p = add_paragraph(doc, line[2:])
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.right_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(55, 86, 35)
        elif re.match(r"^\d+\. \*\*", line):
            clean = re.sub(r"^\d+\. ", "", line)
            p = add_paragraph(doc, clean, style="List Number")
            p.paragraph_format.first_line_indent = Cm(0)
        elif re.match(r"^\d+\. ", line):
            p = add_paragraph(doc, re.sub(r"^\d+\. ", "", line), style="List Number")
            p.paragraph_format.first_line_indent = Cm(0)
        elif line.startswith("- "):
            p = add_paragraph(doc, line[2:], style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
        elif line.startswith(r"\["):
            formula = [line]
            while i + 1 < len(lines) and not lines[i].strip().endswith(r"\]"):
                i += 1
                formula.append(lines[i].strip())
            p = doc.add_paragraph(" ".join(formula).replace("\\[", "").replace("\\]", ""))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for run in p.runs:
                run.font.name = "Cambria Math"
                run.font.size = Pt(10)
        elif line.startswith("**关键词：**"):
            p = add_paragraph(doc, line)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            add_paragraph(doc, line)
        i += 1

    # headers and page numbers
    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("Policy mix - 居民 ABM 方法报告")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(120, 120, 120)
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        footer._p.append(fld)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
